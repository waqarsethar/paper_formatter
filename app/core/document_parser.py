"""Utility functions to identify document structure using python-docx."""

from __future__ import annotations

import re
from docx.document import Document


def strip_heading_number(text: str) -> str:
    """Strip leading number prefixes like '1. ', '1.2 ', '1.2.3 ' from heading text."""
    return re.sub(r"^\d+(?:\.\d+)*[\.\)]\s*", "", text.strip())


# Common reference-section heading variants (lowercase for comparison).
_REFERENCE_HEADINGS = {
    "references",
    "bibliography",
    "works cited",
    "literature cited",
    "citations",
    "reference list",
    "cited literature",
    "literature references",
}


def get_heading_level(paragraph) -> int | None:
    """Return the heading level (1, 2, 3, ...) or None if not a heading.

    Checks ``paragraph.style.name`` for patterns like 'Heading 1',
    'Heading 2', etc.

    Args:
        paragraph: A ``docx.text.paragraph.Paragraph`` instance.

    Returns:
        The heading level as an integer, or ``None`` if the paragraph
        is not styled as a heading.
    """
    style_name = paragraph.style.name or ""
    match = re.match(r"^Heading\s+(\d+)$", style_name)
    if match:
        return int(match.group(1))
    return None


def merge_paragraph_runs(paragraph) -> str:
    """Get the full text from a paragraph by joining all its runs.

    Args:
        paragraph: A ``docx.text.paragraph.Paragraph`` instance.

    Returns:
        The concatenated text of every run in the paragraph.
    """
    return "".join(run.text for run in paragraph.runs)


def is_reference_heading(text: str) -> bool:
    """Check whether *text* matches a common reference-section heading.

    The comparison is case-insensitive and strips leading/trailing
    whitespace.

    Args:
        text: The heading text to test.

    Returns:
        ``True`` if *text* is a recognised reference heading.
    """
    return strip_heading_number(text).lower() in _REFERENCE_HEADINGS


def get_all_sections(doc: Document) -> list[dict]:
    """Return metadata for every headed section in the document.

    Each entry is a dict with:
        - ``heading`` (str): The heading text.
        - ``level`` (int): The heading level (1, 2, 3, ...).
        - ``start`` (int): Paragraph index of the heading itself.
        - ``end`` (int): Paragraph index of the last paragraph in the
          section (exclusive — i.e. the index of the next heading, or
          ``len(doc.paragraphs)`` for the final section).

    Args:
        doc: A ``python-docx`` ``Document`` instance.

    Returns:
        A list of section dicts ordered by their position in the
        document.
    """
    paragraphs = doc.paragraphs
    sections: list[dict] = []

    for idx, para in enumerate(paragraphs):
        level = get_heading_level(para)
        if level is not None:
            sections.append(
                {
                    "heading": para.text,
                    "level": level,
                    "start": idx,
                    "end": len(paragraphs),  # placeholder; patched below
                }
            )

    # Patch each section's end to be the start of the next section.
    for i in range(len(sections) - 1):
        sections[i]["end"] = sections[i + 1]["start"]

    return sections


def find_section_by_heading(
    doc: Document, heading_texts: list[str]
) -> tuple[int, int] | None:
    """Find the paragraph index range for a section identified by heading text.

    The search is **case-insensitive**.  The first heading whose text
    matches any entry in *heading_texts* wins.

    Args:
        doc: A ``python-docx`` ``Document`` instance.
        heading_texts: One or more heading strings to look for.

    Returns:
        A ``(start_idx, end_idx)`` tuple where *start_idx* is the index
        of the heading paragraph and *end_idx* is the index of the next
        heading (or ``len(doc.paragraphs)`` if the section runs to the
        end of the document).  Returns ``None`` if no matching heading
        is found.
    """
    normalised = {t.strip().lower() for t in heading_texts}
    sections = get_all_sections(doc)

    for section in sections:
        heading_text = strip_heading_number(section["heading"]).lower()
        if heading_text in normalised:
            return (section["start"], section["end"])

    return None
