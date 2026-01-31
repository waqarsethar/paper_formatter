"""Citation formatter: detects and converts in-text citation styles."""

from __future__ import annotations

import re
from copy import deepcopy

from docx.text.paragraph import Paragraph

from app.core.document_parser import get_heading_level, merge_paragraph_runs


# ---------------------------------------------------------------------------
# Regex patterns for citation detection
# ---------------------------------------------------------------------------

# APA / author-year: (Smith, 2024), (Smith & Jones, 2024), (Smith et al., 2024b)
# Also matches single entries within multi-citations like (Smith, 2020; Jones, 2019)
AUTHOR_YEAR_RE = re.compile(
    r'\(([A-Z][a-z]+(?:\s(?:&|and)\s[A-Z][a-z]+)*(?:\set\sal\.)?),?\s*(\d{4}[a-z]?)\)'
)

# Multi-citation in one parenthesized group: (Smith, 2020; Jones & Brown, 2019)
MULTI_AUTHOR_YEAR_RE = re.compile(
    r'\(([A-Z][a-z]+(?:\s(?:&|and)\s[A-Z][a-z]+)*(?:\set\sal\.)?),?\s*\d{4}[a-z]?'
    r'(?:\s*;\s*[A-Z][a-z]+(?:\s(?:&|and)\s[A-Z][a-z]+)*(?:\set\sal\.?)?,?\s*\d{4}[a-z]?)+\)'
)

# Individual citation within a multi-citation group
INDIVIDUAL_AUTHOR_YEAR_RE = re.compile(
    r'([A-Z][a-z]+(?:\s(?:&|and)\s[A-Z][a-z]+)*(?:\set\sal\.)?),?\s*(\d{4}[a-z]?)'
)

# Numeric bracket: [1], [1, 2], [1-3], [1; 2; 3]
NUMERIC_BRACKET_RE = re.compile(
    r'\[(\d+(?:\s*[,;\-]\s*\d+)*)\]'
)

# Superscript numbers are detected at the run level (font.superscript),
# so we just need a pattern to recognise bare digit sequences.
SUPERSCRIPT_NUM_RE = re.compile(r'^(\d+(?:\s*[,;\-]\s*\d+)*)$')

STYLE_NAMES = ("author_year", "numeric_bracket", "superscript")


# ---------------------------------------------------------------------------
# Detection helpers
# ---------------------------------------------------------------------------

def _count_superscript_citations(doc, limit: int = 20) -> int:
    """Count paragraphs containing superscript digit runs (up to *limit* body paragraphs)."""
    count = 0
    body_seen = 0
    for para in doc.paragraphs:
        if get_heading_level(para) is not None:
            continue
        body_seen += 1
        if body_seen > limit:
            break
        for run in para.runs:
            if run.font.superscript and SUPERSCRIPT_NUM_RE.search(run.text.strip()):
                count += 1
                break  # one match per paragraph is enough
    return count


def detect_input_style(doc) -> str:
    """Scan the first 20 non-heading paragraphs and return the dominant citation style.

    Returns one of ``"author_year"``, ``"numeric_bracket"``, or ``"superscript"``.
    Falls back to ``"author_year"`` when no citations are detected.
    """
    counts = {"author_year": 0, "numeric_bracket": 0, "superscript": 0}

    body_seen = 0
    for para in doc.paragraphs:
        if get_heading_level(para) is not None:
            continue
        body_seen += 1
        if body_seen > 20:
            break
        text = merge_paragraph_runs(para)
        if AUTHOR_YEAR_RE.search(text) or MULTI_AUTHOR_YEAR_RE.search(text):
            counts["author_year"] += 1
        if NUMERIC_BRACKET_RE.search(text):
            counts["numeric_bracket"] += 1

    counts["superscript"] = _count_superscript_citations(doc, limit=20)

    # Return style with the most matches; break ties by priority order.
    best = max(STYLE_NAMES, key=lambda s: counts[s])
    return best


# ---------------------------------------------------------------------------
# Extraction
# ---------------------------------------------------------------------------

def extract_citations(doc, style: str) -> list[dict]:
    """Return every citation occurrence found in body paragraphs.

    Each dict has:
        - ``para_idx`` (int): paragraph index
        - ``match_text`` (str): the literal text that was matched
        - ``author`` (str | None): author string if available
        - ``year`` (str | None): year string if available
        - ``numbers`` (list[int] | None): numeric ids if available
    """
    citations: list[dict] = []
    for idx, para in enumerate(doc.paragraphs):
        if get_heading_level(para) is not None:
            continue
        text = merge_paragraph_runs(para)

        if style == "author_year":
            # First handle multi-citations like (Smith, 2020; Jones, 2019)
            multi_matches = set()
            for m in MULTI_AUTHOR_YEAR_RE.finditer(text):
                full_match = m.group(0)
                multi_matches.add((m.start(), m.end()))
                # Extract individual citations from within
                inner = full_match[1:-1]  # strip parens
                for im in INDIVIDUAL_AUTHOR_YEAR_RE.finditer(inner):
                    citations.append({
                        "para_idx": idx,
                        "match_text": full_match,
                        "author": im.group(1),
                        "year": im.group(2),
                        "numbers": None,
                        "_is_multi": True,
                    })
            # Then handle single citations, skipping spans already covered
            for m in AUTHOR_YEAR_RE.finditer(text):
                # Skip if this match falls within a multi-citation
                overlaps = any(
                    ms <= m.start() and m.end() <= me
                    for ms, me in multi_matches
                )
                if overlaps:
                    continue
                citations.append({
                    "para_idx": idx,
                    "match_text": m.group(0),
                    "author": m.group(1),
                    "year": m.group(2),
                    "numbers": None,
                })
        elif style == "numeric_bracket":
            for m in NUMERIC_BRACKET_RE.finditer(text):
                nums = _parse_numeric_list(m.group(1))
                citations.append({
                    "para_idx": idx,
                    "match_text": m.group(0),
                    "author": None,
                    "year": None,
                    "numbers": nums,
                })
        elif style == "superscript":
            for run in para.runs:
                if run.font.superscript and SUPERSCRIPT_NUM_RE.search(run.text.strip()):
                    nums = _parse_numeric_list(run.text.strip())
                    citations.append({
                        "para_idx": idx,
                        "match_text": run.text,
                        "author": None,
                        "year": None,
                        "numbers": nums,
                    })
    return citations


def _parse_numeric_list(text: str) -> list[int]:
    """Parse a string like ``'1, 3-5, 7'`` into ``[1, 3, 4, 5, 7]``."""
    nums: list[int] = []
    for part in re.split(r'[,;]\s*', text):
        part = part.strip()
        range_match = re.match(r'(\d+)\s*-\s*(\d+)', part)
        if range_match:
            start, end = int(range_match.group(1)), int(range_match.group(2))
            nums.extend(range(start, end + 1))
        elif part.isdigit():
            nums.append(int(part))
    return nums


# ---------------------------------------------------------------------------
# Citation map
# ---------------------------------------------------------------------------

def build_citation_map(
    citations: list[dict], target_style: str
) -> dict[str, str | int]:
    """Build a mapping from original citation text to a replacement value.

    For *author_year -> numeric* conversions the map assigns sequential
    numbers in order of first appearance.  Other directions may not be
    reliably mappable and will return an empty dict.

    Returns:
        A dict keyed by ``match_text`` whose values are either an ``int``
        (the assigned number) or a ``str`` (replacement text).
    """
    cmap: dict[str, str | int] = {}

    if target_style in ("numeric_bracket", "superscript"):
        # Assign numbers in order of appearance; same author-year pair
        # always gets the same number.
        author_year_counter: dict[str, int] = {}
        current_num = 1
        for cit in citations:
            key = cit["match_text"]
            if key in cmap:
                continue
            # For author-year citations, deduplicate by (author, year).
            ay_key = (cit.get("author"), cit.get("year"))
            if ay_key != (None, None):
                if ay_key not in author_year_counter:
                    author_year_counter[ay_key] = current_num
                    current_num += 1
                cmap[key] = author_year_counter[ay_key]
            else:
                # Numeric sources: keep their existing number(s)
                if cit.get("numbers"):
                    cmap[key] = cit["numbers"][0]
                else:
                    cmap[key] = current_num
                    current_num += 1
    # numeric -> author_year cannot be done without reference metadata.
    return cmap


# ---------------------------------------------------------------------------
# Run-level replacement
# ---------------------------------------------------------------------------

def replace_in_runs(
    paragraph: Paragraph,
    old_text: str,
    new_text: str,
    superscript: bool = False,
) -> bool:
    """Replace *old_text* inside the paragraph's runs with *new_text*.

    The function joins all run texts to locate the character offset of
    *old_text*, then splits or adjusts runs so that the replacement
    occupies its own run (allowing independent formatting).

    If *superscript* is ``True`` the replacement run's
    ``font.superscript`` is set to ``True``.

    Returns ``True`` if a replacement was made, ``False`` otherwise.
    """
    runs = paragraph.runs
    if not runs:
        return False

    full_text = "".join(r.text for r in runs)
    start = full_text.find(old_text)
    if start == -1:
        return False

    end = start + len(old_text)

    # Determine which runs are affected.
    run_boundaries: list[tuple[int, int]] = []  # (start_char, end_char) per run
    offset = 0
    for run in runs:
        run_boundaries.append((offset, offset + len(run.text)))
        offset += len(run.text)

    # Find the first and last run indices that overlap with [start, end).
    first_run_idx: int | None = None
    last_run_idx: int | None = None
    for i, (rs, re_) in enumerate(run_boundaries):
        if rs < end and re_ > start:
            if first_run_idx is None:
                first_run_idx = i
            last_run_idx = i

    if first_run_idx is None or last_run_idx is None:
        return False  # pragma: no cover

    # Collect formatting from the first affected run to preserve styles.
    ref_run = runs[first_run_idx]

    # Build the three text segments: before, replacement, after — all
    # relative to the span of runs we're touching.
    span_start = run_boundaries[first_run_idx][0]
    span_end = run_boundaries[last_run_idx][1]
    span_text = full_text[span_start:span_end]

    before_text = full_text[span_start:start]
    after_text = full_text[end:span_end]

    # Rewrite the first affected run with the before + replacement text.
    # Clear the intermediate runs and put the after text in the last run.
    if first_run_idx == last_run_idx:
        # Everything in one run — split it into up to three runs.
        run = runs[first_run_idx]
        run.text = before_text

        # Insert a new run for the replacement right after.
        new_run = _insert_run_after(paragraph, first_run_idx, new_text, ref_run)
        if superscript:
            new_run.font.superscript = True

        # Insert after-text run if non-empty.
        if after_text:
            _insert_run_after(
                paragraph, first_run_idx + 1, after_text, ref_run
            )
    else:
        # Multiple runs — put before in first, clear middle, after in last.
        runs[first_run_idx].text = before_text

        new_run = _insert_run_after(
            paragraph, first_run_idx, new_text, ref_run
        )
        if superscript:
            new_run.font.superscript = True

        # Clear intermediate runs (offset by 1 because we inserted a run).
        for i in range(first_run_idx + 2, last_run_idx + 2):
            if i < len(paragraph.runs):
                paragraph.runs[i].text = ""

        # Put after-text in the (shifted) last run.
        shifted_last = last_run_idx + 1  # +1 for inserted run
        if shifted_last < len(paragraph.runs):
            paragraph.runs[shifted_last].text = after_text

    return True


def _insert_run_after(paragraph: Paragraph, run_idx: int, text: str, ref_run):
    """Insert a new run with *text* after the run at *run_idx*.

    Copies basic character formatting (font name, size, bold, italic)
    from *ref_run*.

    Returns the newly created run element.
    """
    from docx.oxml.ns import qn
    from copy import deepcopy
    import lxml.etree as etree  # type: ignore[import-untyped]

    runs = paragraph.runs
    ref_element = runs[run_idx]._element

    # Create a new <w:r> element.
    new_r = deepcopy(ref_element)
    # Clear the text node and set our own.
    for t_node in new_r.findall(qn("w:t")):
        new_r.remove(t_node)
    new_t = etree.SubElement(new_r, qn("w:t"))
    new_t.text = text
    # Preserve spaces.
    new_t.set(qn("xml:space"), "preserve")

    # Insert after the reference element.
    ref_element.addnext(new_r)

    # Return the python-docx Run wrapper for the new element.
    from docx.text.run import Run

    new_run = Run(new_r, paragraph)
    return new_run


# ---------------------------------------------------------------------------
# Formatting helpers
# ---------------------------------------------------------------------------

def _format_numeric_bracket(num: int, fmt: str) -> str:
    """Format a number using a template like ``'[{num}]'``."""
    return fmt.format(num=num)


def _format_superscript(num: int, fmt: str) -> str:
    """Format a number for superscript (just the digit text)."""
    return fmt.format(num=num)


# ---------------------------------------------------------------------------
# Main entry-point
# ---------------------------------------------------------------------------

def apply_citations(doc, config: dict) -> dict:
    """Detect in-text citation style and convert to the target style.

    Args:
        doc: A ``python-docx`` ``Document`` instance.
        config: Journal configuration dict.  Must contain
            ``config["citation_style"]`` with keys ``type``, ``format``,
            and ``sort``.

    Returns:
        A dict with ``"warnings"`` (list[str]) and ``"stats"`` (dict).
    """
    warnings: list[str] = []
    stats: dict = {"citations_found": 0, "citations_reformatted": 0}

    cit_config = config.get("citation_style", {})
    target_type = cit_config.get("type", "numeric_bracket")
    target_fmt = cit_config.get("format", "[{num}]")
    target_sort = cit_config.get("sort", "order_of_appearance")

    # Step 1 — detect the input style.
    input_style = detect_input_style(doc)

    # Step 2 — if already the target style, return early.
    if input_style == target_type:
        warnings.append(
            f"Document already uses '{target_type}' citation style; no changes made."
        )
        # Still count citations for stats.
        citations = extract_citations(doc, input_style)
        stats["citations_found"] = len(citations)
        return {"warnings": warnings, "stats": stats}

    # Step 3 — extract citations.
    citations = extract_citations(doc, input_style)
    stats["citations_found"] = len(citations)

    if not citations:
        warnings.append("No citations detected in the document.")
        return {"warnings": warnings, "stats": stats}

    # Step 4 — check feasibility.
    if input_style == "numeric_bracket" and target_type == "author_year":
        warnings.append(
            "Cannot reliably convert numeric citations to author-year format "
            "without reference metadata. Citations left unchanged."
        )
        return {"warnings": warnings, "stats": stats}

    if input_style == "superscript" and target_type == "author_year":
        warnings.append(
            "Cannot reliably convert superscript citations to author-year "
            "format without reference metadata. Citations left unchanged."
        )
        return {"warnings": warnings, "stats": stats}

    # Step 5 — build citation map.
    cit_map = build_citation_map(citations, target_type)

    # Optionally sort alphabetically by author when going from author-year.
    if (
        target_sort == "alphabetical"
        and input_style == "author_year"
        and target_type in ("numeric_bracket", "superscript")
    ):
        # Re-assign numbers alphabetically by (author, year).
        sorted_keys = sorted(
            {(c["author"], c["year"]) for c in citations if c["author"]},
            key=lambda ay: (ay[0].lower(), ay[1]),
        )
        alpha_map = {ay: i + 1 for i, ay in enumerate(sorted_keys)}
        for cit in citations:
            ay = (cit.get("author"), cit.get("year"))
            if ay in alpha_map:
                cit_map[cit["match_text"]] = alpha_map[ay]

    # Step 6 — perform replacements paragraph by paragraph.
    reformatted = 0

    # Group citations by paragraph index for efficient processing.
    by_para: dict[int, list[dict]] = {}
    for cit in citations:
        by_para.setdefault(cit["para_idx"], []).append(cit)

    use_superscript = target_type == "superscript"

    # Track multi-citations already replaced to avoid double-processing.
    replaced_multi: set[tuple[int, str]] = set()

    for para_idx, para_cits in by_para.items():
        para = doc.paragraphs[para_idx]

        # Process citations in reverse order of their position in the
        # paragraph text so that earlier replacements don't shift offsets.
        full_text = merge_paragraph_runs(para)
        para_cits_sorted = sorted(
            para_cits,
            key=lambda c: full_text.rfind(c["match_text"]),
            reverse=True,
        )

        for cit in para_cits_sorted:
            old = cit["match_text"]

            # Handle multi-citations: collect all numbers for the group
            if cit.get("_is_multi"):
                multi_key = (para_idx, old)
                if multi_key in replaced_multi:
                    continue
                replaced_multi.add(multi_key)

                # Gather all numbers for this multi-citation
                multi_nums = []
                for c in para_cits:
                    if c.get("_is_multi") and c["match_text"] == old:
                        n = cit_map.get(c["match_text"])
                        # Try by author+year key
                        ay_key_str = f"({c['author']}, {c['year']})"
                        n = cit_map.get(ay_key_str, n)
                        # Try looking up by individual match
                        for k, v in cit_map.items():
                            if c["author"] and c["author"] in k and c["year"] and c["year"] in k:
                                n = v
                                break
                        if isinstance(n, int):
                            multi_nums.append(n)
                if multi_nums:
                    if use_superscript:
                        new_text = ",".join(str(n) for n in multi_nums)
                    else:
                        nums_str = ", ".join(str(n) for n in multi_nums)
                        new_text = target_fmt.format(num=nums_str)
                    success = replace_in_runs(
                        para, old, new_text, superscript=use_superscript
                    )
                    if success:
                        reformatted += len(multi_nums)
                    else:
                        warnings.append(
                            f"Could not replace multi-citation '{old}'; left unchanged."
                        )
                continue

            num = cit_map.get(old)
            if num is None:
                warnings.append(
                    f"Could not map citation '{old}'; left unchanged."
                )
                continue

            if isinstance(num, int):
                if use_superscript:
                    new_text = _format_superscript(num, target_fmt)
                else:
                    new_text = _format_numeric_bracket(num, target_fmt)
            else:
                new_text = str(num)

            # Handle superscript source runs: the old text lives in a
            # superscript run without surrounding brackets, so we need to
            # also clear the superscript flag when converting away.
            if input_style == "superscript":
                # Find the superscript run directly and replace its text.
                replaced = False
                for run in para.runs:
                    if run.font.superscript and run.text.strip() == old.strip():
                        run.text = new_text
                        if not use_superscript:
                            run.font.superscript = False
                        replaced = True
                        break
                if replaced:
                    reformatted += 1
                else:
                    warnings.append(
                        f"Could not locate superscript run for '{old}'; left unchanged."
                    )
            else:
                success = replace_in_runs(
                    para, old, new_text, superscript=use_superscript
                )
                if success:
                    reformatted += 1
                else:
                    warnings.append(
                        f"Could not replace citation '{old}' at run level; left unchanged."
                    )

    stats["citations_reformatted"] = reformatted

    return {"warnings": warnings, "stats": stats}
