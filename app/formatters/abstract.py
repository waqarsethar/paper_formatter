"""Enhanced abstract formatter: applies formatting to abstract sections."""

from __future__ import annotations

import re

from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


ALIGNMENT_MAP = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
}


def _apply_alignment(paragraph, alignment_name: str) -> None:
    """Set *paragraph* alignment from a string like ``"center"`` or ``"left"``."""
    alignment = ALIGNMENT_MAP.get(alignment_name.lower())
    if alignment is not None:
        paragraph.paragraph_format.alignment = alignment


def _find_abstract_paragraph(doc):
    """Return the paragraph that contains the abstract heading.

    Looks for a paragraph whose text starts with "abstract" (case-insensitive),
    whose style name contains "abstract", or that is a heading paragraph whose
    stripped text equals "abstract".
    """
    for paragraph in doc.paragraphs:
        style_name = (paragraph.style.name or "").lower()
        text = paragraph.text.strip().lower()
        # Strip potential heading number prefix (e.g. "1. Abstract")
        stripped = re.sub(r"^\d+[\.\)]\s*", "", text)
        if stripped.startswith("abstract") or "abstract" in style_name:
            return paragraph
    return None


def _count_abstract_body_words(doc, abstract_para) -> int:
    """Count the words in the abstract body.

    The abstract body is assumed to be the paragraphs between the abstract
    heading and the next headed paragraph (or keywords paragraph).  If the
    abstract heading and body are in the same paragraph (after a colon or
    similar), the body portion of that paragraph is counted instead.
    """
    # Check if abstract text contains body after heading on the same line.
    text = abstract_para.text.strip()
    # Strip the heading portion (e.g. "Abstract" or "ABSTRACT:")
    heading_match = re.match(r"^abstract\s*[:.]?\s*", text, re.IGNORECASE)
    if heading_match:
        remainder = text[heading_match.end():]
        if remainder.strip():
            return len(remainder.split())

    # Otherwise count words in subsequent paragraphs until a heading or
    # keyword line is reached.
    paragraphs = doc.paragraphs
    start_counting = False
    word_count = 0

    # Use text matching to identify the abstract paragraph since object identity doesn't work
    abstract_text = abstract_para.text.strip()

    for para in paragraphs:
        if para.text.strip() == abstract_text and para.style.name == abstract_para.style.name:
            start_counting = True
            continue
        if not start_counting:
            continue

        # Stop at next heading or keywords.
        if para.style.name.startswith("Heading"):
            break
        if para.text.strip().lower().startswith("keyword"):
            break

        word_count += len(para.text.split())

    return word_count


def _get_abstract_body_paragraphs(doc, abstract_para) -> list:
    """Return list of paragraphs in the abstract body (excluding heading)."""
    paragraphs = doc.paragraphs
    start_collecting = False
    body_paras = []

    # Use text matching to identify the abstract paragraph since object identity doesn't work
    abstract_text = abstract_para.text.strip()

    for para in paragraphs:
        if para.text.strip() == abstract_text and para.style.name == abstract_para.style.name:
            start_collecting = True
            continue
        if not start_collecting:
            continue

        # Stop at next heading or keywords.
        if para.style.name.startswith("Heading"):
            break
        if para.text.strip().lower().startswith("keyword"):
            break

        body_paras.append(para)

    return body_paras


def apply_abstract(doc, config) -> dict:
    """Enhanced abstract formatting.

    Reads ``config["abstract"]`` with keys:
        heading_text: str (e.g., "Abstract")
        font_size: float (pt)
        max_words: int (validation)
        alignment: "left" | "center" (optional, default: "left")
        bold_heading: bool (optional, default: True)
        indent_body: float (inches, optional, default: 0.0)
        spacing_after_heading: float (pt, optional, default: 6)

    Args:
        doc: A python-docx ``Document`` object.
        config: Journal configuration dict.

    Returns:
        dict with ``"warnings"`` (list[str]) and ``"stats"`` (dict).
    """
    warnings: list[str] = []

    abstract_config = config.get("abstract", {})
    abstract_para = _find_abstract_paragraph(doc)

    if abstract_para is None:
        warnings.append("Could not identify an abstract section in the document.")
        return {"warnings": warnings, "stats": {"abstract_word_count": 0}}

    # Get configuration options
    heading_text = abstract_config.get("heading_text", "Abstract")
    abstract_font_size = abstract_config.get("font_size", 12)
    max_words = abstract_config.get("max_words")
    alignment = abstract_config.get("alignment", "left")
    bold_heading = abstract_config.get("bold_heading", True)
    indent_body = abstract_config.get("indent_body", 0.0)
    spacing_after_heading = abstract_config.get("spacing_after_heading", 6)

    # Reformat abstract heading text
    current_text = abstract_para.text
    new_text = re.sub(
        r"^abstract\s*[:.]?\s*",
        heading_text + " ",
        current_text,
        count=1,
        flags=re.IGNORECASE,
    ).rstrip()

    # If the replacement ended up only being the heading with trailing
    # space but no body, strip it.
    if new_text.strip() == heading_text:
        new_text = heading_text

    if abstract_para.runs:
        abstract_para.runs[0].text = new_text
        for run in abstract_para.runs[1:]:
            run.text = ""
    else:
        abstract_para.text = new_text

    # Apply font size and bold to the abstract heading paragraph
    for run in abstract_para.runs:
        run.font.size = Pt(abstract_font_size)
        run.font.bold = bold_heading

    # Apply alignment to heading
    _apply_alignment(abstract_para, alignment)

    # Apply spacing after heading
    abstract_para.paragraph_format.space_after = Pt(spacing_after_heading)

    # Format body paragraphs (if separate from heading)
    body_paras = _get_abstract_body_paragraphs(doc, abstract_para)
    for body_para in body_paras:
        # Apply font size
        for run in body_para.runs:
            run.font.size = Pt(abstract_font_size)

        # Apply body indentation
        if indent_body > 0:
            from docx.shared import Inches
            body_para.paragraph_format.left_indent = Inches(indent_body)

    # Word count check
    word_count = _count_abstract_body_words(doc, abstract_para)
    if max_words is not None and word_count > max_words:
        warnings.append(
            f"Abstract contains approximately {word_count} words, "
            f"which exceeds the maximum of {max_words}."
        )

    return {"warnings": warnings, "stats": {"abstract_word_count": word_count}}
